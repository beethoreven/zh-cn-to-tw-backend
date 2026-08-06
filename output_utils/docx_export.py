"""繁化結果文字 <-> .docx 互轉，供下載端點跟「直接上傳」功能使用。"""

from __future__ import annotations

import io

from docx import Document


def text_to_docx_bytes(text: str) -> bytes:
    document = Document()
    # 每一行都當成獨立段落（Word 的「¶」），不是段落內的軟換行
    # （Word 的「↵」）——即使上游正規化過，這裡再切一次雙重保險，
    # 避免殘留的單一換行在 docx 裡變成看不出來的軟換行。
    for line in text.splitlines():
        if line.strip():
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def docx_file_to_text(file_stream) -> str:
    """讀取上傳的 .docx 檔案，逐段落組回文字（雙換行分段，跟本專案
    其他地方的段落慣例一致）。"""
    document = Document(file_stream)
    return "\n\n".join(p.text for p in document.paragraphs)
